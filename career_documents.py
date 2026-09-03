from pathlib import Path
from typing import Dict, List, Tuple
import io, re
from pypdf import PdfReader
from docx import Document

SUPPORTED={".pdf",".docx",".txt",".md",".csv",".json",".log"}

def clean_text(text:str)->str:
    text=text.replace("\x00"," "); text=re.sub(r"[ \t]+"," ",text); text=re.sub(r"\n{3,}","\n\n",text)
    return text.strip()

def extract_document(name:str,data:bytes)->Tuple[str,Dict]:
    ext=Path(name).suffix.lower()
    if ext not in SUPPORTED: raise ValueError(f"Unsupported file type: {ext}")
    if ext==".pdf":
        reader=PdfReader(io.BytesIO(data)); parts=[f"[Page {i}]\n{p.extract_text() or ''}" for i,p in enumerate(reader.pages,1)]; text="\n\n".join(parts); meta={"pages":len(reader.pages),"type":"PDF"}
    elif ext==".docx":
        doc=Document(io.BytesIO(data)); parts=[p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            parts.extend(" | ".join(c.text.strip() for c in row.cells) for row in table.rows)
        text="\n".join(parts); meta={"pages":None,"type":"DOCX"}
    else:
        text=data.decode("utf-8",errors="ignore"); meta={"pages":None,"type":ext.upper().lstrip(".")}
    text=clean_text(text)
    if not text: raise ValueError(f"No readable text found in {name}")
    return text,meta

def classify_document(name:str,text:str)->str:
    sample=(name+"\n"+text[:6000]).lower()
    rules=[("Resume",["curriculum vitae","resume","professional summary","work experience","career profile"]),("Position Description",["position description","job description","key accountabilities","about the role"]),("Cover Letter",["cover letter","dear hiring manager"]),("Performance Review",["performance review","performance appraisal","objectives and key results"]),("Certification",["certification","certificate of completion","credential"])]
    for label,needles in rules:
        if any(n in sample for n in needles): return label
    return "Other"

def chunk_text(text:str,chunk_size:int=1200,overlap:int=180)->List[str]:
    if chunk_size<=0 or overlap<0 or overlap>=chunk_size: raise ValueError("chunk_size must be positive and overlap must be smaller than chunk_size")
    words=text.split(); step=chunk_size-overlap; chunks=[]
    for start in range(0,len(words),step):
        chunk=" ".join(words[start:start+chunk_size])
        if chunk: chunks.append(chunk)
        if start+chunk_size>=len(words): break
    return chunks
