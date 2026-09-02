"""Document ingestion with provenance for PDF, DOCX, and scanned PDFs."""

from __future__ import annotations

from io import BytesIO
from pathlib import Path
from typing import List

from career_intelligence import DocumentChunk


class DocumentIngestionError(Exception):
    """Raised when a document cannot be read safely."""


def ingest_pdf(pdf_bytes: bytes, source: str = "uploaded PDF",
               ocr_enabled: bool = True) -> List[DocumentChunk]:
    try:
        from pypdf import PdfReader
        reader = PdfReader(BytesIO(pdf_bytes))
    except Exception as error:
        raise DocumentIngestionError(f"Unable to read PDF: {error}") from error

    chunks = []
    for page_number, page in enumerate(reader.pages, start=1):
        text = (page.extract_text() or "").strip()
        if not text and ocr_enabled:
            text = _ocr_page(pdf_bytes, page_number - 1)
        if text:
            chunks.append(DocumentChunk(text=text, source=source, page=page_number,
                                        metadata={"format": "pdf", "page": str(page_number)}))
    return chunks


def ingest_docx(docx_bytes: bytes, source: str = "uploaded DOCX") -> List[DocumentChunk]:
    try:
        from docx import Document
        document = Document(BytesIO(docx_bytes))
    except ImportError as error:
        raise DocumentIngestionError(
            "DOCX support requires python-docx. Install project requirements.") from error
    except Exception as error:
        raise DocumentIngestionError(
            f"Unable to read DOCX: {error}") from error

    paragraphs = [paragraph.text.strip()
                  for paragraph in document.paragraphs if paragraph.text.strip()]
    if not paragraphs:
        raise DocumentIngestionError(
            "The DOCX contains no readable paragraphs.")
    return [DocumentChunk("\n\n".join(paragraphs), source=source,
                          metadata={"format": "docx"})]


def ingest_document(content: bytes, filename: str, ocr_enabled: bool = True) -> List[DocumentChunk]:
    suffix = Path(filename).suffix.lower()
    if suffix == ".pdf":
        return ingest_pdf(content, filename, ocr_enabled)
    if suffix == ".docx":
        return ingest_docx(content, filename)
    raise DocumentIngestionError("Only PDF and DOCX documents are supported.")


def _ocr_page(pdf_bytes: bytes, page_index: int) -> str:
    try:
        import pypdfium2 as pdfium
        import pytesseract
        document = pdfium.PdfDocument(pdf_bytes)
        page = document[page_index]
        text = pytesseract.image_to_string(
            page.render(scale=2.2).to_pil()).strip()
        page.close()
        document.close()
        return text
    except ImportError as error:
        raise DocumentIngestionError(
            "This scanned PDF needs pypdfium2, pytesseract, and the Tesseract binary.") from error
    except Exception as error:
        raise DocumentIngestionError(
            f"OCR failed on page {page_index + 1}: {error}") from error
